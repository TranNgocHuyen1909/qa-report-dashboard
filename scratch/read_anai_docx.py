import sys
import docx

def read_docx(file_path, output_path):
    doc = docx.Document(file_path)
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("# CONTENTS OF ANAI.docx\n\n")
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                f.write(f"[{i}] [{para.style.name}] {para.text}\n\n")
        
        f.write("\n\n# TABLES\n\n")
        for ti, table in enumerate(doc.tables):
            f.write(f"## Table {ti}: {len(table.rows)} rows x {len(table.columns)} cols\n\n")
            for ri, row in enumerate(table.rows):
                cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                f.write(f"- Row {ri}: {cells}\n")

if __name__ == '__main__':
    read_docx(r'd:\ANAI\qa-report-dashboard\ANAI.docx', r'd:\ANAI\qa-report-dashboard\scratch\docx_anai_content.md')
